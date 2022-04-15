import re
from datetime import datetime
import openpyxl
from docx import Document



wk = openpyxl.load_workbook(r"C:\Users\shill\PycharmProjects\Excel\TestWritePyExcel.xlsx")


sheets = wk.sheetnames

#print(sheets[0])


sh = wk[sheets[0]]

Speaker1 = sh['B3'].value

list_val = []

list_rep = []

first_row = sh.min_row
last_rows = sh.max_row
max_columns = sh.max_column

ranger = range(first_row, last_rows + 1)

def list_maker(list, column,):
    for i in ranger:
        #print(i)
        output = sh[column + str(i)].value
        if output == 'TODAY':
            output = datetime.today().strftime('%m/%d/%Y')
        elif column == 'A':
            output = str('<<' + output + '>>')
        else:
            output = output

        list.append(output)


        #print(str(i) + '. ' + str(output))






def docx_replace_regex(doc_obj, regex_list, replace_list):
    for (regex, replace) in zip(regex_list, replace_list):
        regexs = re.compile(r'{}'.format(str(regex)))
        replaces = r'{}'.format(replace)

        print('regex {} : replace {}'.format(regex, replace))

        for p in doc_obj.paragraphs:
            #print('{}'.format(p))
            if regexs.search(p.text):
                inline = p.runs
                print('{}'.format(inline))
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regexs.search(inline[i].text):
                        text = regexs.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    print('{}:{}:{}'.format(cell, regex, replace))
                    docx_replace_regex(cell, regex, replace)



list_maker(list_val, 'B')
list_maker(list_rep, 'A')
list_val1 = 'Spencer Hill'
list_rep1 = '<<Speaker1>>'
#regex1 = re.compile(r'<Speaker1>')
#replace1 = r'{}'.format(Speaker1)
#print(list_val)
#print('-----------------')
#print(list_rep)

filename = "test.docx"
doc = Document(filename)
docx_replace_regex(doc, list_val1, list_rep1)
doc.save('result1.docx')